# # rag_utils.py
# import os
# import logging
# import numpy as np
# import openai
# from faster_whisper import WhisperModel
# from moviepy import AudioFileClip
# import PyPDF2
# import docx
# import pandas as pd
# import pytesseract
# from pytesseract import TesseractNotFoundError
# from dotenv import load_dotenv

# load_dotenv()
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # OpenAI API key
# openai.api_key = os.getenv("OPENAI_API_KEY")

# # Constants
# UPLOAD_DIR    = "uploads"
# EMBED_MODEL   = "text-embedding-ada-002"
# CHUNK_SIZE    = 500
# CHUNK_OVERLAP = 50

# # Whisper for audio/video
# _whisper = WhisperModel("base", device="cpu", compute_type="int8")


# def transcribe_and_cache_audio(path: str) -> str:
#     """
#     Transcribe audio/video once, cache to `<path>.txt`, and return the transcript.
#     """
#     txt_path = f"{path}.txt"
#     # If we’ve already transcribed, load it
#     if os.path.exists(txt_path):
#         try:
#             return open(txt_path, "r", encoding="utf8").read()
#         except Exception as e:
#             logger.warning("Could not read cached transcript %s: %s", txt_path, e)

#     # Otherwise, transcribe anew
#     try:
#         # Handle video → audio
#         if path.lower().endswith((".mp4", ".mov", ".mkv")):
#             wav = f"{path}.wav"
#             AudioFileClip(path).write_audiofile(wav, verbose=False, logger=None)
#             segments, _ = _whisper.transcribe(wav)
#         else:
#             segments, _ = _whisper.transcribe(path)

#         text = " ".join(seg.text for seg in segments)

#         # Cache it
#         try:
#             with open(txt_path, "w", encoding="utf8") as f:
#                 f.write(text)
#         except Exception as e:
#             logger.warning("Failed to write transcript cache %s: %s", txt_path, e)

#         return text
#     except Exception as e:
#         logger.error("Transcription failed for %s: %s", path, e)
#         return ""


# def list_uploaded_files() -> list[str]:
#     return [
#         os.path.join(UPLOAD_DIR, fn)
#         for fn in os.listdir(UPLOAD_DIR)
#         if os.path.isfile(os.path.join(UPLOAD_DIR, fn))
#     ]


# def extract_text(path: str) -> str:
#     """
#     Extract readable text from any supported file.
#     Audio/video → cached whisper transcript.
#     Images, PDF, docx, CSV/XLSX, TXT → respective libraries.
#     """
#     suffix = os.path.splitext(path)[1].lower()
#     try:
#         # Audio/video
#         if suffix in (".mp3", ".wav", ".m4a", ".mp4", ".mov", ".mkv"):
#             return transcribe_and_cache_audio(path)

#         # Images
#         if suffix in (".jpg", ".jpeg", ".png"):
#             try:
#                 return pytesseract.image_to_string(path)
#             except TesseractNotFoundError:
#                 logger.warning("Tesseract not installed; skipping OCR for %s", path)
#                 return ""

#         # PDFs
#         if suffix == ".pdf":
#             reader = PyPDF2.PdfReader(path)
#             return "\n".join(page.extract_text() or "" for page in reader.pages)

#         # Word docs
#         if suffix == ".docx":
#             doc = docx.Document(path)
#             return "\n".join(p.text for p in doc.paragraphs)

#         # Plain text
#         if suffix == ".txt":
#             return open(path, encoding="utf8").read()

#         # CSV / Excel
#         if suffix == ".csv":
#             return pd.read_csv(path).to_string(index=False)
#         if suffix == ".xlsx":
#             return pd.read_excel(path, engine="openpyxl").to_string(index=False)

#         return ""
#     except Exception as e:
#         logger.error("extract_text error for %s: %s", path, e)
#         return ""


# def chunk_text(text: str):
#     """Yield overlapping word‑based chunks."""
#     words = text.split()
#     step = CHUNK_SIZE - CHUNK_OVERLAP
#     for i in range(0, len(words), step):
#         yield " ".join(words[i : i + CHUNK_SIZE])


# def embed_text(text: str) -> list[float]:
#     """Embed a single chunk via OpenAI embeddings."""
#     resp = openai.embeddings.create(model=EMBED_MODEL, input=[text])
#     return resp.data[0].embedding


# def retrieve_similar_chunks(query: str, top_k: int = 3):
#     """
#     1) If the user asked for the ‘first sentence’, find the first audio/video file,
#        pull its cached transcript, and immediately return [(1.0, path, 1, first_sentence)].
#     2) Otherwise, run RAG over everything in uploads/.
#     """
#     # 1) Special-case “first sentence” requests
#     if "first sentence" in query.lower():
#         for path in list_uploaded_files():
#             if os.path.splitext(path)[1].lower() in (".mp3", ".wav", ".m4a", ".mp4", ".mov", ".mkv"):
#                 text = transcribe_and_cache_audio(path)
#                 first = text.split(".")[0].strip() + "."
#                 return [(1.0, path, 1, first)]
#         return []

#     # 2) Normal RAG
#     qv = np.array(embed_text(query))
#     qn = np.linalg.norm(qv)
#     if qn == 0:
#         return []

#     corpus = []
#     for path in list_uploaded_files():
#         txt = extract_text(path)
#         for idx, chunk in enumerate(chunk_text(txt), start=1):
#             try:
#                 corpus.append((embed_text(chunk), path, idx, chunk))
#             except Exception as e:
#                 logger.warning("Skipping chunk %s[%d]: %s", path, idx, e)

#     sims = []
#     for emb, path, idx, chunk in corpus:
#         ev = np.array(emb)
#         denom = np.linalg.norm(ev) * qn
#         if denom > 0:
#             sims.append((float(np.dot(ev, qv) / denom), path, idx, chunk))

#     sims.sort(key=lambda x: x[0], reverse=True)
#     return sims[:top_k]





















# rag_utils.py
import os
import logging
import numpy as np
import openai

from faster_whisper import WhisperModel
from moviepy import AudioFileClip
import PyPDF2
import docx
import pandas as pd
import pytesseract
from pytesseract import TesseractNotFoundError
from dotenv import load_dotenv

from PIL import Image, ImageOps

# Optional fallback OCR
try:
    import easyocr
except ImportError:
    easyocr = None

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Constants
UPLOAD_DIR    = "uploads"
EMBED_MODEL   = "text-embedding-ada-002"
CHUNK_SIZE    = 500
CHUNK_OVERLAP = 50

# Whisper for audio/video
_whisper = WhisperModel("base", device="cpu", compute_type="int8")


def transcribe_and_cache_audio(path: str) -> str:
    """
    Transcribe audio/video once, cache to `<path>.txt`, and return the transcript.
    """
    txt_path = f"{path}.txt"
    if os.path.exists(txt_path):
        try:
            return open(txt_path, "r", encoding="utf8").read()
        except Exception as e:
            logger.warning("Could not read cached transcript %s: %s", txt_path, e)

    try:
        if path.lower().endswith((".mp4", ".mov", ".mkv")):
            wav = f"{path}.wav"
            AudioFileClip(path).write_audiofile(wav, verbose=False, logger=None)
            segments, _ = _whisper.transcribe(wav)
        else:
            segments, _ = _whisper.transcribe(path)

        text = " ".join(seg.text for seg in segments)

        try:
            with open(txt_path, "w", encoding="utf8") as f:
                f.write(text)
        except Exception as e:
            logger.warning("Failed to write transcript cache %s: %s", txt_path, e)

        return text
    except Exception as e:
        logger.error("Transcription failed for %s: %s", path, e)
        return ""


def list_uploaded_files() -> list[str]:
    return [
        os.path.join(UPLOAD_DIR, fn)
        for fn in os.listdir(UPLOAD_DIR)
        if os.path.isfile(os.path.join(UPLOAD_DIR, fn))
    ]


def extract_text(path: str) -> str:
    """
    Extract readable text from any supported file.
    - Audio/video → cached Whisper transcript
    - Images → Tesseract (pre‑processed) or EasyOCR
    - PDF → PyPDF2
    - DOCX → python‑docx
    - CSV / XLSX / TXT → pandas & plain open()
    """
    suffix = os.path.splitext(path)[1].lower()
    try:
        # ─── Audio/video ─────────────────────────────────────
        if suffix in (".mp3", ".wav", ".m4a", ".mp4", ".mov", ".mkv"):
            return transcribe_and_cache_audio(path)

        # ─── Image ───────────────────────────────────────────
        if suffix in (".jpg", ".jpeg", ".png"):
            # 1) Pre‑process for Tesseract
            try:
                img = Image.open(path)
                img = ImageOps.grayscale(img)
                w, h = img.size
                img = img.resize((w * 2, h * 2), Image.Resampling.LANCZOS)
                text = pytesseract.image_to_string(img, config="--psm 6")
                if text.strip():
                    return text
            except TesseractNotFoundError:
                logger.warning("Tesseract not installed; skipping OCR for %s", path)
            except Exception as e:
                logger.error("Error during image OCR for %s: %s", path, e)

            # 2) Fallback to EasyOCR if available
            if easyocr:
                reader = easyocr.Reader(["en"], gpu=False)
                return "\n".join(reader.readtext(path, detail=0))

            return ""

        # ─── PDF ──────────────────────────────────────────────
        if suffix == ".pdf":
            reader = PyPDF2.PdfReader(path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        # ─── DOCX ────────────────────────────────────────────
        if suffix == ".docx":
            d = docx.Document(path)
            return "\n".join(p.text for p in d.paragraphs)

        # ─── Plain text ─────────────────────────────────────
        if suffix == ".txt":
            return open(path, encoding="utf8").read()

        # ─── CSV / XLSX ─────────────────────────────────────
        if suffix == ".csv":
            return pd.read_csv(path).to_string(index=False)
        if suffix == ".xlsx":
            return pd.read_excel(path, engine="openpyxl").to_string(index=False)

        return ""
    except Exception as e:
        logger.error("extract_text error for %s: %s", path, e)
        return ""


def chunk_text(text: str):
    """Yield overlapping word‑based chunks of size CHUNK_SIZE."""
    words = text.split()
    step = CHUNK_SIZE - CHUNK_OVERLAP
    for i in range(0, len(words), step):
        yield " ".join(words[i : i + CHUNK_SIZE])


def embed_text(text: str) -> list[float]:
    """Embed a single chunk via OpenAI embeddings API."""
    resp = openai.embeddings.create(model=EMBED_MODEL, input=[text])
    return resp.data[0].embedding


def retrieve_similar_chunks(query: str, top_k: int = 3):
    """
    1) If the user asks for "first sentence", return it directly from the first audio/video file.
    2) Otherwise, standard RAG over all chunks in uploads/.
    """
    # Special‑case “first sentence” requests
    if "first sentence" in query.lower():
        for path in list_uploaded_files():
            ext = os.path.splitext(path)[1].lower()
            if ext in (".mp3", ".wav", ".m4a", ".mp4", ".mov", ".mkv"):
                text = transcribe_and_cache_audio(path)
                first = text.split(".")[0].strip() + "."
                return [(1.0, path, 1, first)]
        return []

    # Normal RAG
    qv = np.array(embed_text(query))
    qn = np.linalg.norm(qv)
    if qn == 0:
        return []

    corpus = []
    for path in list_uploaded_files():
        txt = extract_text(path)
        for idx, chunk in enumerate(chunk_text(txt), start=1):
            try:
                corpus.append((embed_text(chunk), path, idx, chunk))
            except Exception as e:
                logger.warning("Skipping chunk %s[%d]: %s", path, idx, e)

    sims = []
    for emb, path, idx, chunk in corpus:
        ev = np.array(emb)
        denom = np.linalg.norm(ev) * qn
        if denom > 0:
            sims.append((float(np.dot(ev, qv) / denom), path, idx, chunk))

    sims.sort(key=lambda x: x[0], reverse=True)
    return sims[:top_k]